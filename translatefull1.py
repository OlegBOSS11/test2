import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import pyttsx3
import threading
from deep_translator import GoogleTranslator
from docx import Document
import PyPDF2
from fpdf import FPDF

# Переместите это в начало вашего скрипта после импортов
translator_instance = GoogleTranslator()
languages = translator_instance.get_supported_languages()

# Инициализация TTS движка
engine = pyttsx3.init()


def split_text(text, max_length=5000):
    parts = []
    while text:
        part = text[:max_length]
        first_newline = part.rfind('\n')
        if first_newline != -1 and len(part) > max_length / 2:
            parts.append(part[:first_newline])
            text = text[first_newline + 1:]
        else:
            parts.append(part)
            text = text[max_length:]
    return parts


def translate_text(text, src_lang='auto', dest_lang='ru'):
    translator_instance1 = GoogleTranslator(source=src_lang, target=dest_lang)
    parts = split_text(text)
    translated_text = ""
    for part in parts:
        translated_text += translator_instance1.translate(part) + "\n"
    return translated_text


def perform_translation():
    source_text = input_text.get("1.0", "end-1c")
    src_lang = src_lang_var.get()
    dest_lang = dest_lang_var.get()
    translated_text = translate_text(source_text, src_lang, dest_lang)
    output_text.delete("1.0", "end")
    output_text.insert("1.0", translated_text)


def process_file(file_path, src_lang, dest_lang):
    _, ext = os.path.splitext(file_path)
    if ext.lower() in ['.txt']:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    elif ext.lower() in ['.pdf']:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = "".join(page.extract_text() or '' for page in reader.pages)
    elif ext.lower() in ['.docx']:
        doc = Document(file_path)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)
    else:
        messagebox.showerror("Ошибка", "Неподдерживаемый формат файла")
        return None
    return translate_text(text, src_lang, dest_lang)


def save_translated_text(translated_text, output_path):
    _, ext = os.path.splitext(output_path)
    if ext.lower() in ['.txt']:
        with open(output_path, 'w', encoding='utf-8') as file:
            file.write(translated_text)
    elif ext.lower() in ['.pdf']:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, translated_text)
        pdf.output(output_path)
    elif ext.lower() in ['.docx']:
        doc = Document()
        for paragraph in translated_text.split('\n'):
            doc.add_paragraph(paragraph)
        doc.save(output_path)
    else:
        messagebox.showerror("Ошибка", "Неподдерживаемый формат файла для сохранения")


def translate_file():
    # Эта функция запрашивает у пользователя файл, переводит его и предлагает сохранить результат.
    file_path = filedialog.askopenfilename()
    if not file_path:
        return
    src_lang = src_lang_var.get()
    dest_lang = dest_lang_var.get()
    translated_text = process_file(file_path, src_lang, dest_lang)
    if translated_text:
        save_translated_file(translated_text)


def save_translated_file(translated_text):
    # Эта функция сохраняет переведенный текст в файл без его отображения в поле вывода.
    output_path = filedialog.asksaveasfilename(defaultextension=".txt")
    if output_path:
        save_translated_text(translated_text, output_path)
        messagebox.showinfo("Успех", f"Переведенный файл сохранен в: {output_path}")


def save_translated_text_to_file():
    translated_text = output_text.get("1.0", "end-1c")  # Извлекаем переведенный текст
    if not translated_text.strip():
        messagebox.showinfo("Пусто", "Нет переведенного текста для сохранения.")
        return
    output_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt")])
    if output_path:  # Проверка, что пользователь выбрал файл
        with open(output_path, 'w', encoding='utf-8') as file:
            file.write(translated_text)
        messagebox.showinfo("Успех", f"Переведенный файл сохранен в: {output_path}")


def speak_text(text):
    def run():
        engine.say(text)
        engine.runAndWait()

    threading.Thread(target=run).start()


def stop_speaking():
    engine.stop()


def clear_fields():
    input_text.delete("1.0", "end")
    output_text.delete("1.0", "end")


def copy_text_from_input():
    try:
        root.clipboard_clear()
        root.clipboard_append(input_text.get(tk.SEL_FIRST, tk.SEL_LAST))
    except tk.TclError:
        pass


def paste_text_to_input():
    try:
        input_text.insert(tk.INSERT, root.clipboard_get())
    except tk.TclError:
        pass


def copy_text_from_output():
    try:
        root.clipboard_clear()
        root.clipboard_append(output_text.get(tk.SEL_FIRST, tk.SEL_LAST))
    except tk.TclError:
        pass


def create_context_menu():
    input_menu = tk.Menu(root, tearoff=0)
    input_menu.add_command(label="Копировать", command=copy_text_from_input)
    input_menu.add_command(label="Вставить", command=paste_text_to_input)

    output_menu = tk.Menu(root, tearoff=0)
    output_menu.add_command(label="Копировать", command=copy_text_from_output)

    def show_input_menu(event):
        input_menu.tk_popup(event.x_root, event.y_root)

    def show_output_menu(event):
        output_menu.tk_popup(event.x_root, event.y_root)

    input_text.bind("<Button-3>", show_input_menu)
    output_text.bind("<Button-3>", show_output_menu)


root = tk.Tk()
root.title("Переводчик файлов")
root.minsize(400, 300)  # Установите минимальный размер окна

frame = ttk.Frame(root, padding="10")
frame.grid(sticky="WENS")

root.grid_columnconfigure(0, weight=1)  # Добавьте вес для главного столбца окна
root.grid_rowconfigure(0, weight=1)  # Добавьте вес для главной строки окна

# Добавьте вес для строк и столбцов в frame, чтобы элементы растягивались
# frame.grid_rowconfigure(1, weight=1)
# frame.grid_columnconfigure(1, weight=1)
frame.grid_rowconfigure(0, weight=1)
frame.grid_rowconfigure(2, weight=1)
frame.grid_columnconfigure(0, weight=1)
frame.grid_columnconfigure(2, weight=1)

input_text = tk.Text(frame, height=10, width=50)
input_text.grid(row=0, column=0, columnspan=3, sticky="WENS")

# Для входного текста
input_text_scroll = tk.Scrollbar(frame, orient="vertical", command=input_text.yview)
input_text_scroll.grid(row=0, column=3, sticky="ns")
input_text.configure(yscrollcommand=input_text_scroll.set)

output_text = tk.Text(frame, height=10, width=50)
output_text.grid(row=2, column=0, columnspan=3, sticky="WENS")
# Для выходного текста
output_text_scroll = tk.Scrollbar(frame, orient="vertical", command=output_text.yview)
output_text_scroll.grid(row=2, column=3, sticky="ns")
output_text.configure(yscrollcommand=output_text_scroll.set)

# Инициализация переменных для хранения выбранного языка
src_lang_var = tk.StringVar(value='auto')
dest_lang_var = tk.StringVar(value='ru')

# Настройка веса строк и столбцов
root.grid_rowconfigure(0, weight=1)
root.grid_columnconfigure(0, weight=1)
frame.grid_rowconfigure(0, weight=1)
frame.grid_columnconfigure(0, weight=1)

src_lang_menu = ttk.Combobox(frame, textvariable=src_lang_var, values=languages, state="readonly")
src_lang_menu.grid(row=1, column=0, sticky="EW")
dest_lang_menu = ttk.Combobox(frame, textvariable=dest_lang_var, values=languages, state="readonly")
dest_lang_menu.grid(row=1, column=2, sticky="EW")

translate_button = ttk.Button(frame, text="Перевести", command=translate_file)
translate_button.grid(row=3, column=0, sticky="EW")

# Кнопки
open_file_button = ttk.Button(frame, text="Открыть файл и перевести ", command=translate_file)
open_file_button.grid(row=3, column=0, sticky="EW")

translate_button = ttk.Button(frame, text="Перевести", command=perform_translation)
translate_button.grid(row=3, column=1, sticky="EW")

save_button = ttk.Button(frame, text="Сохранить перевод", command=save_translated_text_to_file)
save_button.grid(row=3, column=2, sticky="EW")
speak_input_button = ttk.Button(frame, text="Озвучить оригинал", command=lambda: speak_text(input_text.get("1.0",
                                                                                                           "end-1c")))
speak_input_button.grid(row=4, column=0, sticky="EW")

speak_output_button = ttk.Button(frame, text="Озвучить перевод", command=lambda: speak_text(output_text.get("1.0",
                                                                                                            "end-1c")))
speak_output_button.grid(row=4, column=2, sticky="EW")

stop_button = ttk.Button(frame, text="Остановить", command=stop_speaking)
stop_button.grid(row=5, column=0, sticky="EW")

clear_button = ttk.Button(frame, text="Очистить поля", command=clear_fields)
clear_button.grid(row=5, column=2, sticky="EW")

create_context_menu()

root.mainloop()
